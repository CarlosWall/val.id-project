pip install Flask

from flask import Flask, request, render_template
from werkzeug.utils import secure_filename
import pandas as pd
import numpy as np
import seaborn as sns
import matplotlib.pyplot as plt
from scipy.stats import linregress
from docx import Document
from io import BytesIO
from docx.shared import Inches

app = Flask(__name__)

def process_data(file):
    data = pd.read_excel(file)
    dados = data['Dados']
    
    # Calcular métricas estatísticas
    mean = dados.mean()
    median = dados.median()
    mode = dados.mode()[0]
    std_dev = dados.std()
    quartiles = np.percentile(dados, [25, 50, 75])
    iqr = quartiles[2] - quartiles[0]
    lower_bound = quartiles[0] - 1.5 * iqr
    upper_bound = quartiles[2] + 1.5 * iqr
    outliers = dados[(dados < lower_bound) | (dados > upper_bound)]
    cv = std_dev / mean * 100
    bias = median - mean
    percentile_2_5 = np.percentile(dados, 2.5)
    percentile_97_5 = np.percentile(dados, 97.5)
    
    # Gerar gráficos
    plt.figure(figsize=(12, 6))
    plt.subplot(1, 2, 1)
    sns.boxplot(data=dados)
    plt.title('Boxplot')
    boxplot_bytes = BytesIO()
    plt.savefig(boxplot_bytes, format='png')
    plt.close()
    
    plt.figure(figsize=(12, 6))
    plt.subplot(1, 2, 1)
    sns.histplot(data=dados, kde=True)
    plt.title('Histograma')
    histogram_bytes = BytesIO()
    plt.savefig(histogram_bytes, format='png')
    plt.close()
    
    plt.figure(figsize=(8, 6))
    x = np.arange(len(dados))
    slope, intercept, _, _, _ = linregress(x, dados)
    trendline = slope * x + intercept
    plt.plot(x, dados, label='Dados')
    plt.plot(x, trendline, label='Regressão Linear', color='red')
    plt.title('Regressão Linear')
    plt.legend()
    regression_bytes = BytesIO()
    plt.savefig(regression_bytes, format='png')
    plt.close()
    
    # Criar arquivo DOC com os resultados e gráficos
    doc = Document()
    doc.add_heading('Análise Estatística', level=1)
    doc.add_paragraph(f'Média: {mean}')
    doc.add_paragraph(f'Mediana: {median}')
    doc.add_paragraph(f'Moda: {mode}')
    doc.add_paragraph(f'Desvio Padrão: {std_dev}')
    doc.add_paragraph(f'Coeficiente de Variação: {cv}')
    doc.add_paragraph(f'Bias: {bias}')
    doc.add_paragraph(f'Quartis: Q1={quartiles[0]}, Q2={quartiles[1]}, Q3={quartiles[2]}')
    doc.add_paragraph(f'Outliers: {outliers.tolist()}')
    doc.add_paragraph(f'Limite Inferior: {lower_bound}')
    doc.add_paragraph(f'Limite Superior: {upper_bound}')
    doc.add_paragraph(f'Percentil 2.5: {percentile_2_5}')
    doc.add_paragraph(f'Percentil 97.5: {percentile_97_5}')
    
    doc.add_heading('Gráficos', level=2)
    doc.add_picture(boxplot_bytes, width=Inches(6))
    doc.add_picture(histogram_bytes, width=Inches(6))
    doc.add_picture(regression_bytes, width=Inches(6))
    
    result = BytesIO()
    doc.save(result)
    result.seek(0)
    
    return result.read()

@app.route('/')
def index():
    return render_template('index.html')

@app.route('/upload', methods=['POST'])
def upload():
    file = request.files['file']
    if file.filename == '':
        return 'Nenhum arquivo selecionado.'
    
    if file:
        filename = secure_filename(file.filename)
        file.save(filename)
        result = process_data(filename)
        return result, 200, {'Content-Type': 'application/msword'}
    else:
        return 'Erro no upload do arquivo.', 400

if __name__ == '__main__':
    app.run(debug=True)
